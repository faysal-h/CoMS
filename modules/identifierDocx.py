import logging


from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm
from docx.enum.style import WD_STYLE_TYPE


logger = logging.getLogger('CCMS.IdentifierDocx')


class IdentifiersDocument():
    def __init__(self) -> None:
        self.document = Document()
        self.createTwoColumnsPage()

    def createTwoColumnsPage(self):
        section = self.document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

    def add_styles(self):
        styles = self.document.styles

        style1 = styles.add_style('Bold16', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = styles["Normal"]
        fontOfStyle1 = style1.font
        fontOfStyle1.name = "Times New Roman"
        fontOfStyle1.size = Pt(16)
        fontOfStyle1.bold = True
        paragraphFormat = style1.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

        style2 = styles.add_style('Bold12', WD_STYLE_TYPE.PARAGRAPH)
        style2.base_style = styles["Normal"]
        fontOfStyle2 = style2.font
        fontOfStyle2.name = "Times New Roman"
        fontOfStyle2.size = Pt(12)
        fontOfStyle2.bold = True
        paragraphFormat = style2.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

    def PageLayout(self, size):
        self.size = size
        if self.size == "A4":
            sections = self.document.sections
            sectionMain = sections[0]
            # Page dimension and header footer distance
            sectionMain.page_height = Mm(297)
            sectionMain.page_width = Mm(210)
            sectionMain.top_margin = Inches(0.5)
            sectionMain.bottom_margin = Inches(0.5)
            sectionMain.left_margin = Inches(0.5)
            sectionMain.right_margin = Inches(0.5)
            sectionMain.header_distance = Inches(0.0)
            sectionMain.footer_distance = Inches(0.0)

            return 'First Section of A4 pages size is created.'
        else:
            return 'Page size not supported.'

    def addHeader(self, HeaderText:str):
        section = self.document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = str(HeaderText)
        #header.is_linked_to_previous

    def saveDoc(self, saveLocation, IdentifiersORevnelops='Identifiers'):

        self.document.save(saveLocation)
        logger.info(f"Identififers file saved in {saveLocation}")
        return saveLocation

    def addFileIdentifiers(self, caseNo1, caseNo2, parcels, fir, ps, district, BatchDate=""):
        id = self.document.add_paragraph("", style="Bold16")
        id_format = id.paragraph_format
        id_format.space_after = Pt(0)
        id.add_run(f'\t\t\t({BatchDate})\n').font.size = Pt(8)
        id.add_run('Case No(s) :\t').font.size = Pt(11)
        id.add_run(f'{caseNo1}\n').font.size = Pt(12)
        if caseNo2 not in [None, '']:
            id.add_run(f'\t\t{caseNo2}\n').font.size = Pt(11)
        id.add_run(f'Parcels:\t{parcels}\n').font.size = Pt(11)
        id.add_run(f'FIR:\t\t{fir}\n').font.size = Pt(11)
        id.add_run(f'PS: \t\t{ps}\n').font.size = Pt(11)
        id.add_run(f'District:\t{district}\n').font.size = Pt(11)
        id.add_run('').font.size = Pt(11)

    def addEnvelopsIdentifiers(self, caseNo1, AddressTo, district):
        id = self.document.add_paragraph("", style="Bold16")
        id_format = id.paragraph_format
        id_format.space_after = Pt(0)
        id.add_run(f'\n\n\t{caseNo1}\n').font.size = Pt(10)
        id.add_run('To:\n').font.size = Pt(11)
        id.add_run(f'\t{AddressTo},\n').font.size = Pt(13)
        id.add_run(f'\t{district}.\n').font.size = Pt(13)
        id.add_run('').font.size = Pt(11)

class NotesDocument(IdentifiersDocument):
    def __init__(self) -> None:
        self.document = Document()

    def addNote(self, caseNo1, caseNo2, parcels):
        id = self.document.add_paragraph("", style="Bold16")
        id_format = id.paragraph_format
        id_format.space_after = Pt(0)
        id.add_run('_______________________________________________________________________________________________\n').font.size = Pt(11)
        id.add_run('Case No(s):\t').font.size = Pt(11)
        id.add_run(f'{caseNo1}\t{caseNo2}\n').font.size = Pt(12)
        id.add_run(f'Total Parcels:\t{parcels}\n').font.size = Pt(10)
        # id.add_run('').font.size = Pt(10)

    def addParcelDetailsInNotes(self, parcelNo, itemNo, quantity, caliber, itemDetail):
        id = self.document.add_paragraph("")
        id_format = id.paragraph_format
        id_format.space_after = Pt(0)
        id.add_run(f'Parcel {parcelNo}: ').font.size = Pt(10)
        id.add_run(f'{quantity} ({itemNo}) {caliber}').font.size = Pt(10)
        id.add_run('').font.size = Pt(10)


if __name__ == '__main__':

    pass
    # i = IdentifiersDocument()
    # i.PageLayout('A4')
    # i.add_styles()
    # i.createTwoColumnsPage()
    # # i.tableIdentifiersFiles("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    # i.addFileIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", parcels=6
    #                     ,fir=6, firDate="02.02.2022", ps='abc', district='xyz')
    # i.addEnvelopsIdentifiers(caseNo1="PFSA2020-123456-FTM-123456", AddressTo="CPO", district="Pakpattan")
    # i.saveDoc()
