import logging

import inflect

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_LINE_SPACING


logging.basicConfig(level=logging.DEBUG)



'''
NOTE There should be a template docx in the folder with the following custom sytles
BulletCustomNormal , type= Paragraph, bullets enabled
TableGridCustom ,   type = Table

NOTE The page numbering field should also be enabled as page numbering is not supported at this moment.
'''
class Report():
    def __init__(self):
        self.document = Document('./modules/templates/template.docx')

    def testFiresStatementFromItemNo(self, EvType: str, itemNo: str):
        itemsToCheck = {
                        'R': f': test fires produced in the lab {itemNo}TC1 & {itemNo}TC2',
                        'P': f': test fires produced in the lab {itemNo}TC1 & {itemNo}TC2',
                        'S': f': test fires produced in the lab {itemNo}TS1 & {itemNo}TS2',
                        'M': f': test fires produced in the lab {itemNo}TC1 & {itemNo}TC2',
                        }
        logging.info(f"Evidence Type is {EvType}, Item No is {itemNo}")

        if itemNo == None or "":
            return ""
        else:
            itemNo = itemNo.upper()
            if(EvType == "firearm"):
                # loops through dictionary to find respective item test fires names
                for key in itemsToCheck:
                    if(itemNo.find(key) != -1):
                        logging.info(f"this test fire is found {itemsToCheck[key]}")
                        return itemsToCheck[key]
            else:
                return ""
    
    def header(self, caseNo):
        section0 = self.document.sections[0]

        # Set Different header and footer for first page
        # section0.different_first_page_header_footer = True

        # # # Footer of FIRST PAGE
        # first_page_header = section0.first_page_header
        
        # paragraph = first_page_header.paragraphs[0]
        # paragraph.text = 'First Page Header'

        
        # Default Header from Second page to onward pages
        header = section0.header
        paragraph = header.paragraphs[0]
        paragraph.text = f"\n\n\n\n\n\n\n\n\n\t\tCase#: {caseNo}\n"

    def accusedStatementfrmName(self, accusedName):
            if(accusedName not in [None, '']):
                return f"\n(said to be recovered from the accused {accusedName})"
            else:
                return ""

    def addRowInTableEvidenceDetails(self, parcelNumber:int, submissionDate:str,
                                    submitterName:str, submitterRank:str, fir:str, firDate:str,
                                    PS:str, District:str, quantityInWords:str, caliber:str, EVDetails:str,
                                    itemString:str, itemNumbers:str,testFires:str,accused:str ):
        
        tableMain = self.document.tables[1]

        newRowCells = tableMain.add_row().cells

        # set alignment of each cell to top in Row
        for cell in newRowCells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Parcel NUMBER CELL
        newRowCells[0].paragraphs[0].add_run(f'{parcelNumber}',style='SimpleText')
        
        # SUBMITTER CELL
        newRowCells[1].paragraphs[0].add_run(f'{submitterName} ({submitterRank}) \n{submissionDate}',
                                             style='SimpleText')
        
        # FIR & PS CELL
        newRowCells[2].paragraphs[0].add_run(f'{fir}/{firDate[8:]},'
                        f' \n{PS}, {District}', style='SimpleText')

        # ITEM DETAILS CELL
        newRowCells[3].paragraphs[0].add_run(f'{quantityInWords} {caliber} caliber {EVDetails} '
                        f'({itemString} {itemNumbers}{testFires}){accused}', style='SimpleText')


    #NOTE THIS FUNCTION CREATE AND STORE CUSTOM STYLE
    def add_styles(self):
        styles = self.document.styles
        
        style1 = styles.add_style('Bold', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = styles["Normal"]
        fontOfStyle1 = style1.font
        fontOfStyle1.name = "Times New Roman"
        fontOfStyle1.size = Pt(11)
        fontOfStyle1.bold = True
        paragraphFormat = style1.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

        style2 = styles.add_style('BoldUnderline', WD_STYLE_TYPE.PARAGRAPH)
        style2.base_style = styles["Normal"]
        fontOfStyle2 = style2.font
        fontOfStyle2.name = "Times New Roman"
        fontOfStyle2.size = Pt(14)
        fontOfStyle2.bold = True
        fontOfStyle2.underline = True
        paragraphFormat = style2.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1

        style = styles.add_style('CompactParagraph', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        fontOfStyle = style.font
        fontOfStyle.name = "Times New Roman"
        fontOfStyle.size = Pt(12)
        #fontOfStyle.bold = True
        #fontOfStyle.underline = True
        paragraphFormat = style.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1
        #paragraphFormat.left_indent = Mm(2)
        
        style4 = styles.add_style('TableHeading', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('SimpleText', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('TableStyle', WD_STYLE_TYPE.TABLE)
        style4.base_style = styles["Normal Table"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        return print('Custom Styles added to the word self.document.')

    #CREATE A SECTION AND SET MARGINS OF IT
    def PageLayout(self, size):
        self.size = size
        if self.size == "A4":
            sections = self.document.sections
            sectionMain = sections[0]
            # Page dimension and header footer distance
            sectionMain.page_height = Mm(297)
            sectionMain.page_width = Mm(210)
            sectionMain.top_margin = Inches(2.14)
            sectionMain.bottom_margin = Inches(0.87)
            sectionMain.left_margin = Inches(0.75)
            sectionMain.right_margin = Inches(0.7)
            sectionMain.header_distance = Inches(.39)
            # sectionMain.footer_distance = Inches(1.18)

            return 'First Section of A4 pages size is created.'
        else:
            return 'Page size not supported.'
    
    #CREATE HEADING OF THE REPORT
    def paraTOD(self):
        titleOfDocument = self.document.add_paragraph("", style="Bold")
        titleOfDocument.add_run("Firearms & Toolmarks Examination Report").font.size = Pt(12)
        titleOfDocument_format = titleOfDocument.paragraph_format
        titleOfDocument_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titleOfDocument_format.space_before = Pt(0)
        titleOfDocument_format.space_after = Pt(0)

    #CASE NUMBER TABLE
    def tableCaseDetails(self, caseNo1, caseNo2, addressee, district):
        
        tableCaseDetails = self.document.tables[0]
        #TABLE STYLE
        #tableCaseDetails.columns[0].width = Cm(1)
        # tableCaseDetails.style = 'TableGridCustom'
        # tableCaseDetails.allow_autofit =False
        
        #Length of table is 6309360
        # tableCaseDetails.rows[0].cells[0].width = Mm(32)
        # tableCaseDetails.rows[0].cells[1].width = Mm(70)
        # tableCaseDetails.rows[0].cells[2].width = Mm(32)
        # tableCaseDetails.rows[0].cells[3].width = Mm(52)
        # tableCaseDetails.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # tableCaseDetails.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # tableCaseDetails.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # tableCaseDetails.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        #TABLE VALUES
        firstRowCells = tableCaseDetails.rows[1].cells
        # firstRowCells[0].paragraphs[0].add_run('Agency Case#',style='TableHeading')
        firstRowCells[1].paragraphs[0].add_run(f'{caseNo1}', style='SimpleText')
        if(caseNo2 in [None, "", "None"] ):
            pass
        else:
            firstRowCells[1].paragraphs[0].add_run(f'\n{caseNo2}', style='SimpleText')
        # firstRowCells[2].paragraphs[0].add_run('Attention To', style='TableHeading')
        firstRowCells[3].paragraphs[0].add_run(f'{addressee}, {district}.', style='SimpleText')

    def paraEvDetail(self, Addressee, District, items, testRequest):
        # if(items>1):
        #     wasORwere = "items were"
        # else:
        #     wasORwere = "item was"
            
        if(testRequest == None or testRequest == ''):
            testRequest = "Comparison of Cartridge Cases and Shotshell Cases with Submitted Firearms and Functionality Testing"

        #NOTE EVIDENCE SUBMISSION PARAGRAPH
        # evidenceDetailsParagraph = self.document.add_paragraph("", style='CompactParagraph')
        # evidenceDetailsParagraph_format = evidenceDetailsParagraph.paragraph_format
        # evidenceDetailsParagraph_format.space_before = Pt(2)
        # evidenceDetailsParagraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        EVdescriptionParagraph = self.document.paragraphs[2]
        EVdescriptionParagraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        EVdescriptionParagraph.add_run(f" submitted along with the request of {Addressee}, {District} for ",
                                         style='SimpleText')
        EVdescriptionParagraph.add_run(f"{testRequest}.").bold =True
        # evidenceDetailsParagraph2 = self.document.add_paragraph("", style='CompactParagraph')

    #CREATE TABLE OF EVIDENCE INFORMATION
    def tableEvDetails(self, parcels):

        tableEVDetails = self.document.tables[1]
        # tableEVDetails.style = 'TableGridCustom'
        # tableEVDetails.allow_autofit = False
     

        # firstRowCells = tableEVDetails.rows[0].cells
        # firstRowCells[0].paragraphs[0].add_run('Parcel#',style='TableHeading')
        # firstRowCells[1].paragraphs[0].add_run('Submitter &\nSubmission Date', style='TableHeading')
        # firstRowCells[2].paragraphs[0].add_run('FIR & PS', style='TableHeading')
        # firstRowCells[3].paragraphs[0].add_run('Evidence Details\nItem No#', style='TableHeading')
        
        for i, parcel in enumerate(parcels, start=0):

            # converts quantity of items from digits to words and other variables to statements
            quantityInWords = inflect.engine().number_to_words(parcel[10])
            accused = self.accusedStatementfrmName(parcel[14])
            testFires = self.testFiresStatementFromItemNo(EvType=parcel[7], itemNo=parcel[9])
            itemsOrItems = 'Item' if parcel[10] < 2 else "Items" 


            if(i==0):
                # for first entry in list of parcels first row must be created otherwise it will be added to heading
                self.addRowInTableEvidenceDetails(  parcelNumber=parcel[0],
                                                    submissionDate=parcel[1],
                                                    submitterName=parcel[2],
                                                    submitterRank=parcel[3],
                                                    fir=parcel[4],
                                                    firDate=parcel[5],
                                                    PS=parcel[12],
                                                    District=parcel[13],
                                                    quantityInWords=quantityInWords,
                                                    caliber=parcel[6],
                                                    EVDetails=parcel[8],
                                                    itemString=itemsOrItems,
                                                    itemNumbers=parcel[9],
                                                    testFires=testFires,
                                                    accused=accused
                                                    )

            else:
                #  NOTE parcels[i-1][0] Previous Parcel Number.
                #  As -1 points to last item of list, so this also works for first Parcel of list
                if(parcel[0] != parcels[i-1][0]):
                    
                    # for first entry in list of parcels first row must be created otherwise it will be added to heading
                    self.addRowInTableEvidenceDetails(  parcelNumber=parcel[0],
                                                        submissionDate=parcel[1],
                                                        submitterName=parcel[2],
                                                        submitterRank=parcel[3],
                                                        fir=parcel[4],
                                                        firDate=parcel[5],
                                                        PS=parcel[12],
                                                        District=parcel[13],
                                                        quantityInWords=quantityInWords,
                                                        caliber=parcel[6],
                                                        EVDetails=parcel[8],
                                                        itemString=itemsOrItems,
                                                        itemNumbers=parcel[9],
                                                        testFires=testFires,
                                                        accused=accused
                                                    )

                else:
                    # move to last row of table
                    previousRowCells = tableEVDetails.rows[-1].cells

                    previousRowCells[3].paragraphs[0].add_run(f' and {quantityInWords} {parcel[6]} caliber {parcel[8]} '
                                    f'({itemsOrItems} {parcel[9]}{testFires}){accused}', style='SimpleText')

        # # Column 1 PARCEL NO WIDTH
        # for cell in tableEVDetails.columns[0].cells:
        #     cell.width = Mm(10)

        # # Column 2 WIDTH
        # for cell in tableEVDetails.columns[1].cells:
        #     cell.width = Mm(35)
        
        # # Column 4 WIDTH
        # for cell in tableEVDetails.columns[2].cells:
        #     cell.width = Mm(35)

        # # Column 4 WIDTH
        # for cell in tableEVDetails.columns[3].cells:
        #     cell.width = Mm(90)

        #This is to seprate next table from this one
        seprationPara = self.document.add_paragraph(style='CompactParagraph')
        seprationPara.add_run('   ').font.size = Pt(1)
        paragraph_format = seprationPara.paragraph_format
        seprationPara.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(2)

    #CREATE TABLE OF ANALYSIS INFORMATION
    def tableAnalysisDetails(self, startDate : str, endDate : str):
        tableAnalysis = self.document.tables[2]
        # tableAnalysis.style = 'TableGridCustom'
        # tableAnalysis.allow_autofit = False
        #Length of table is 180mm
        # tableAnalysis.rows[0].cells[0].width = Mm(10)
        # tableAnalysis.rows[0].cells[1].width = Mm(50)
        # tableAnalysis.rows[0].cells[2].width = Mm(90)
        # tableAnalysis.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # tableAnalysis.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # tableAnalysis.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # firstRowCells = tableAnalysis.rows[0].cells
        # firstRowCells[0].paragraphs[0].add_run('Analysis Start Date',style='TableHeading')
        # firstRowCells[1].paragraphs[0].add_run('Analysis Completion Date', style='TableHeading')
        # firstRowCells[2].paragraphs[0].add_run('Examination Method/ Tests Performed', style='TableHeading')

        secondRowCells = tableAnalysis.rows[1].cells
        secondRowCells[0].paragraphs[0].add_run(f'{startDate}')
        secondRowCells[1].paragraphs[0].add_run(f'{endDate}')
        # secondRowCells[2].paragraphs[0].add_run('Physical Examination, Comparison Microscopy, Test Firing and ABIS Scanning',
        #                                         style='SimpleText')

        # adjust column 1 length
        for cell in tableAnalysis.columns[0].cells:
            cell.width = Mm(38)

        # adjust column 2 length
        for cell in tableAnalysis.columns[1].cells:
            cell.width = Mm(48)

    #CREATE CONCLUSION, Have to remove space after paragraph
    def paraResults(self):
        resultsHeading = self.document.add_paragraph("", style="BoldUnderline")
        resultsHeading_format = resultsHeading.paragraph_format
        resultsHeading_format.space_after = Pt(0)
        resultsHeading.add_run('Details of Results and Conclusions Based on Test(s) Performed:').font.size = Pt(11)

        
        for i in listResults:
            self.document.add_paragraph(f"{i}",style="BulletCustomNormal")


    #CREATE NOTE(S)
    def paraNotes(self):
        notesHeading = self.document.add_paragraph("", style="BoldItalic")
        notesHeading.add_run(f'Note(s):').font.size = Pt(11)

        # listNotes = ['This is the first note.', 'This is the second note.']
        # for i in listNotes:
        #     self.document.add_paragraph(style="BulletCustomNormal").add_run(f"{i}", style="SimpleText").font.italic = True

    #CREATE DIPOSITION OF EVIDENCE PARAGRAPH
    def paraDisposition(self):
        dispositionHeading = self.document.add_paragraph("", style="BoldUnderline")
        dispositionHeading.add_run('Disposition of Evidence:').font.size = Pt(11)
        dispositionParagraph = self.document.add_paragraph(f'', style='CompactParagraph')
        dispositionParagraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        EORParagraph = self.document.add_paragraph('', style='Bold')
        EORParagraph.add_run('X...End of Report...X').font.size = Pt(12)
        EORParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def footer(self):
        # FOOTER ANALYST AND REVIEWER TODO NEEDS TO MODIFY TO ADD FOOTER VARIABLE
        sections = self.document.sections
        sectionMain = sections[0]
        footers = sectionMain.footer
        paragraphFooter = footers.paragraphs[0]
        paragraphFooter.text = ""
        for run in paragraphFooter.runs:
            run.font.size = Pt(10)

    def save(self, saveLocation):
        self.document.save(saveLocation)

if __name__ == '__main__':
    testReport = Report()
    testReport.PageLayout('A4')
    # testReport.add_styles()
    testReport.paraTOD()
    testReport.tableCaseDetails()
    testReport.paraEvDetail()
    testReport.tableEvDetails()
    testReport.tableAnalysisDetails()
    testReport.paraResults()
    testReport.paraNotes()
    testReport.paraDisposition()
    testReport.footer()
    testReport.save()

    # os.system("start ./TestReport.docx")