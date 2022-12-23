import logging

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from sqlalchemy import table

from modules.CusPath import UserPaths

logging.basicConfig(level=logging.DEBUG)


'''
NOTE There should be a template docx in the folder with the following custom sytles
BulletCustomNormal , type= Paragraph, bullets enabled
TableGridCustom ,   type = Table
NOTE The page numbering field should also be enabled as page numbering is not supported at this moment.
'''


class CPRDocument():
    def __init__(self):
        self.document = Document(UserPaths.cprTemplatePath)

    # NOTE THIS FUNCTION CREATE AND STORE CUSTOM STYLE

    def add_styles(self):
        styles = self.document.styles

        style1 = styles.add_style('Bold', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = styles["Normal"]
        fontOfStyle1 = style1.font
        fontOfStyle1.name = "Times New Roman"
        fontOfStyle1.size = Pt(11)
        fontOfStyle1.bold = True
        paragraphFormat = style1.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

        style2 = styles.add_style('BoldUnderline', WD_STYLE_TYPE.PARAGRAPH)
        style2.base_style = styles["Normal"]
        fontOfStyle2 = style2.font
        fontOfStyle2.name = "Times New Roman"
        fontOfStyle2.size = Pt(14)
        fontOfStyle2.bold = True
        fontOfStyle2.underline = True
        paragraphFormat = style2.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1

        style = styles.add_style('CompactParagraph', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        fontOfStyle = style.font
        fontOfStyle.name = "Times New Roman"
        fontOfStyle.size = Pt(12)
        #fontOfStyle.bold = True
        #fontOfStyle.underline = True
        paragraphFormat = style.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1
        #paragraphFormat.left_indent = Mm(2)

        style4 = styles.add_style('TableHeading', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('SimpleText', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('TableStyle', WD_STYLE_TYPE.TABLE)
        style4.base_style = styles["Normal Table"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        return print('Custom Styles added to the word self.document.')

    # CREATE A SECTION AND SET MARGINS OF IT
    def PageLayout(self, size):
        self.size = size
        if self.size == "A4":
            sections = self.document.sections
            sectionMain = sections[0]
            # Page dimension and header footer distance
            sectionMain.page_height = Mm(297)
            sectionMain.page_width = Mm(210)
            sectionMain.top_margin = Inches(2.14)
            sectionMain.bottom_margin = Inches(0.87)
            sectionMain.left_margin = Inches(0.75)
            sectionMain.right_margin = Inches(0.7)
            sectionMain.header_distance = Inches(.39)
            # sectionMain.footer_distance = Inches(1.18)

            return 'First Section of A4 pages size is created.'
        else:
            return 'Page size not supported.'

    # CASE NUMBER TABLE
    def addRowInMainTable(self, Serial: int, CaseNo: str, FIR: str, PS: str, District: str):

        tableMain = self.document.tables[0]

        newRow = tableMain.add_row().cells

        newRow[0].paragraphs[0].add_run(Serial)
        newRow[1].paragraphs[0].add_run(CaseNo)
        newRow[4].paragraphs[0].add_run(f'{FIR}')
        newRow[5].paragraphs[0].add_run(f'{PS.title()}, {District.title()}')

    def save(self, saveLocation):
        self.document.save(saveLocation)


if __name__ == '__main__':
    testSheet = CPRDocument()

    table = testSheet.document.tables[0]
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(f'123456')
    # testReport.PageLayout('A4')
    # testReport.add_styles()
    # testSheet.addRowInMainTable(Serial=1, CaseNo='ABC123',
    #                             FIR='123/12', PS='Test', District='xyz')
    testSheet.save()

    # os.system("start ./TestReport.docx")
