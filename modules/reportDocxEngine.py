import os

from docx import Document
from docx.shared import Inches, Pt, Mm, Emu
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

'''
NOTE There should be a template docx in the folder with the following custom sytles
BulletCustomNormal , type= Paragraph, bullets enabled
TableGridCustom ,   type = Table
NOTE The page numbering field should also be enabled as page numbering is not supported at this moment.
'''
class Reports():
    def __init__(self):
        self.document = Document('./Word/template.docx')
        
    #NOTE THIS FUNCTION CREATE AND STORE CUSTOM STYLE
    def add_styles(self):
        styles = self.document.styles
        
        style1 = styles.add_style('Bold', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = styles["Normal"]
        fontOfStyle1 = style1.font
        fontOfStyle1.name = "Times New Roman"
        fontOfStyle1.size = Pt(14)
        fontOfStyle1.bold = True
        paragraphFormat = style1.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

        style2 = styles.add_style('BoldUnderline', WD_STYLE_TYPE.PARAGRAPH)
        style2.base_style = styles["Normal"]
        fontOfStyle2 = style2.font
        fontOfStyle2.name = "Times New Roman"
        fontOfStyle2.size = Pt(14)
        fontOfStyle2.bold = True
        fontOfStyle2.underline = True
        paragraphFormat = style2.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1

        style = styles.add_style('CompactParagraph', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        fontOfStyle = style.font
        fontOfStyle.name = "Times New Roman"
        fontOfStyle.size = Pt(12)
        #fontOfStyle.bold = True
        #fontOfStyle.underline = True
        paragraphFormat = style.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        paragraphFormat.line_spacing = 1
        #paragraphFormat.left_indent = Mm(2)
        
        style4 = styles.add_style('TableHeading', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('SimpleText', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style4 = styles.add_style('TableStyle', WD_STYLE_TYPE.TABLE)
        style4.base_style = styles["Normal Table"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        return print('Custom Styles added to the word self.document.')

    #CREATE A SECTION AND SET MARGINS OF IT
    def PageLayout(self, size):
        self.size = size
        if self.size == "A4":
            sections = self.document.sections
            sectionMain = sections[0]
            # Page dimension and header footer distance
            sectionMain.page_height = Mm(297)
            sectionMain.page_width = Mm(210)
            sectionMain.top_margin = Inches(2.14)
            sectionMain.bottom_margin = Inches(0.87)
            sectionMain.left_margin = Inches(0.75)
            sectionMain.right_margin = Inches(0.7)
            sectionMain.header_distance = Inches(.39)
            sectionMain.footer_distance = Inches(1.18)

            return 'First Section of A4 pages size is created.'
        else:
            return 'Page size not supported.'
    
    
    #CREATE HEADING OF THE REPORT
    def paraTOD(self):
        titleOfDocument = self.document.add_paragraph(f"Firearms & Toolmarks Examination Report", style="Bold")
        titleOfDocument_format = titleOfDocument.paragraph_format
        titleOfDocument_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titleOfDocument_format.space_before = Pt(0)
        titleOfDocument_format.space_after = Pt(0)

    #CASE NUMBER TABLE
    def tableCaseDetails(self):
        
        tableCaseDetails = self.document.add_table(rows=1, cols=4)
        #TABLE STYLE
        #tableCaseDetails.columns[0].width = Cm(1)
        tableCaseDetails.style = 'TableGridCustom'
        tableCaseDetails.allow_autofit =False
        #Length of table is 6309360
        tableCaseDetails.rows[0].cells[0].width = Mm(30)
        tableCaseDetails.rows[0].cells[1].width = Mm(70)
        tableCaseDetails.rows[0].cells[2].width = Mm(28)
        tableCaseDetails.rows[0].cells[3].width = Mm(52)
        # tableCaseDetails.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # tableCaseDetails.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # tableCaseDetails.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # tableCaseDetails.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #TABLE VALUES
        firstRowCells = tableCaseDetails.rows[0].cells
        firstRowCells[0].paragraphs[0].add_run('Agency Case#',style='TableHeading')
        firstRowCells[1].paragraphs[0].add_run('PFSA20XX-XXXXXX-FTM-XXXXXX', style='SimpleText')
        firstRowCells[2].paragraphs[0].add_run('Attention To:', style='TableHeading')
        firstRowCells[3].paragraphs[0].add_run('SP Investigation, Cantt Division, Lahore.', style='SimpleText')

    def paraEvDetail(self):
        #NOTE EVIDENCE SUBMISSION PARAGRAPH
        evidenceDetailsParagraph = self.document.add_paragraph("", style='CompactParagraph')
        evidenceDetailsParagraph_format = evidenceDetailsParagraph.paragraph_format
        evidenceDetailsParagraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        evidenceDetailsHeading = evidenceDetailsParagraph.add_run("Description of Evidence Submitted\n", style='SimpleText')
        evidenceDetailsHeading.bold = True
        evidenceDetailsHeading.underline = True
        evidenceDetailsParagraph.add_run("The following evidence item was submitted on 02.01.2017 by Muhammad Sarwar (ASI) along with the request of DPO, Vehari for ", style='SimpleText')
        evidenceDetailsParagraph.add_run(f"Comparison of Bullet and Functionality Testing.").bold =True

    #CREATE TABLE OF EVIDENCE INFORMATION
    def tableEvDetails(self):
        
        tableEVDetails = self.document.add_table(rows=1, cols=4)
        tableEVDetails.style = 'TableGridCustom'
        tableEVDetails.allow_autofit =False
        #NOTE length of table is 180mm
        tableEVDetails.rows[0].cells[0].width = Mm(15)
        tableEVDetails.rows[0].cells[1].width = Mm(50)
        tableEVDetails.rows[0].cells[2].width = Mm(53)
        tableEVDetails.rows[0].cells[3].width = Mm(62)
        tableEVDetails.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableEVDetails.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableEVDetails.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableEVDetails.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        firstRowCells = tableEVDetails.rows[0].cells
        firstRowCells[0].paragraphs[0].add_run('Parcel#',style='TableHeading')
        firstRowCells[1].paragraphs[0].add_run('Submitter &\nSubmission Date', style='TableHeading')
        firstRowCells[2].paragraphs[0].add_run('FIR & PS', style='TableHeading')
        firstRowCells[3].paragraphs[0].add_run('Evidence Details\nItem No#', style='TableHeading')
        #This is to seprate next table from this one
        self.document.add_paragraph('', style='CompactParagraph')

    #CREATE TABLE OF ANALYSIS INFORMATION
    def tableAnalysisDetails(self):
        tableAnalysis = self.document.add_table(rows=1, cols=3)
        tableAnalysis.style = 'TableGridCustom'
        tableAnalysis.allow_autofit = False
        #Length of table is 180mm
        tableAnalysis.rows[0].cells[0].width = Mm(40)
        tableAnalysis.rows[0].cells[1].width = Mm(50)
        tableAnalysis.rows[0].cells[2].width = Mm(90)
        tableAnalysis.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableAnalysis.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableAnalysis.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        firstRowCells = tableAnalysis.rows[0].cells
        firstRowCells[0].paragraphs[0].add_run('Analysis Start Date',style='TableHeading')
        firstRowCells[1].paragraphs[0].add_run('Analysis Completion Date', style='TableHeading')
        firstRowCells[2].paragraphs[0].add_run('Examination Method/ Tests Performed', style='TableHeading')

    #CREATE CONCLUSION, Have to remove space after paragraph
    def paraResults(self):
        resultsHeading = self.document.add_paragraph("", style="BoldUnderline")
        resultsHeading_format = resultsHeading.paragraph_format
        resultsHeading_format.space_after = Pt(0)
        resultsHeading.add_run('Details of Results and Conclusions Based on Test(s) Performed:').font.size = Pt(12)

        listResults = ['This is the first result.', 'This is the second result.', 'This is the third result.']
        for i in listResults:
            self.document.add_paragraph(f"{i}",style="BulletCustomNormal")


    #CREATE NOTE(S)
    def paraNotes(self):
        notesHeading = self.document.add_paragraph("", style="BoldUnderline")
        notesHeading.add_run('Note(s):').font.size = Pt(12)
        listNotes = ['This is the first note.', 'This is the second note.']
        for i in listNotes:
            self.document.add_paragraph(style="BulletCustomNormal").add_run(f"{i}", style="SimpleText").font.italic = True

    #CREATE DIPOSITION OF EVIDENCE PARAGRAPH
    def paraDisposition(self):
        dispositionHeading = self.document.add_paragraph("", style="BoldUnderline")
        dispositionHeading.add_run('Disposition of Heading:').font.size = Pt(12)
        dispositionParagraph = self.document.add_paragraph('The case property/ evidence may be received by the responsible official of your office on submitting authorization letter/docket within 15 days after the receipt of this report.  Ammunition components should be maintained for possible future examinations.', style='CompactParagraph')
        dispositionParagraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        EORParagraph = self.document.add_paragraph('', style='Bold')
        EORParagraph.add_run('X...End of Report...X').font.size = Pt(12)
        EORParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def footer(self):
        # FOOTER ANALYST AND REVIEWER TODO NEEDS TO MODIFY TO ADD FOOTER VARIABLE
        sections = self.document.sections
        sectionMain = sections[0]
        footers = sectionMain.footer
        paragraphFooter = footers.paragraphs[0]
        paragraphFooter.text = "\tThis is footer"

    def save(self):
        self.document.save("./Word/TestReport.docx")

if __name__ == '__main__':
    testReport = Reports()
    testReport.PageLayout('A4')
    # testReport.add_styles()
    testReport.paraTOD()
    testReport.tableCaseDetails()
    testReport.paraEvDetail()
    testReport.tableEvDetails()
    testReport.tableAnalysisDetails()
    testReport.paraResults()
    testReport.paraNotes()
    testReport.paraDisposition()
    testReport.footer()
    testReport.save()

    os.system("start ./Word/TestReport.docx")