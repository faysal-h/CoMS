import os

from CusPath import UserPaths

from docx import Document
from docx.shared import Inches, Pt, Mm, Emu
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

templateIdentifier = os.path.join(os.getcwd(), "modules", "templates", "Identifiers.docx")
# templateIdentifier.replace("\\", "/")
class Reports():
    def __init__(self):
        self.document = Document(templateIdentifier)
        
        #NOTE THIS FUNCTION CREATE AND STORE CUSTOM STYLE
    def add_styles(self):
        styles = self.document.styles
        
        style1 = styles.add_style('Bold10', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = styles["Normal"]
        fontOfStyle1 = style1.font
        fontOfStyle1.name = "Times New Roman"
        fontOfStyle1.size = Pt(10)
        fontOfStyle1.bold = True
        paragraphFormat = style1.paragraph_format
        paragraphFormat.space_before = Pt(0)
        paragraphFormat.space_after = Pt(0)
        #fontOfStyle1.underline = True

        #NOTE These are style used for runs in tables.
        style4 = styles.add_style('TableHeading', WD_STYLE_TYPE.CHARACTER)
        style4.base_style = styles["Normal"]
        fontOfStyle4 = style4.font
        fontOfStyle4.name = "Times New Roman"
        fontOfStyle4.size = Pt(16)
        fontOfStyle4.bold = True
        #fontOfStyle3.underline = True

        style5 = styles.add_style('TableStyleBody', WD_STYLE_TYPE.CHARACTER)
        style5.base_style = styles["Normal Table"]
        fontOfStyle5 = style5.font
        fontOfStyle5.name = "Times New Roman"
        fontOfStyle5.size = Pt(12)
        fontOfStyle5.bold = True

        # Table Grid Design
        style7 = styles.add_style('TableStyle', WD_STYLE_TYPE.TABLE)
        style7.base_style = styles["Light Grid"]
        fontOfStyle7 = style7.font
        fontOfStyle7.name = "Times New Roman"
        #fontOfStyle3.size = Pt(12)
        #fontOfStyle4.bold = True
        #fontOfStyle3.underline = True
        
        return print('Custom Styles added to the word self.document.')
    
    def PageLayout(self, size):
        self.size = size
        if self.size == "A4":
            sections = self.document.sections
            sectionMain = sections[0]
            # Page dimension and header footer distance
            sectionMain.page_height = Mm(297)
            sectionMain.page_width = Mm(210)
            sectionMain.top_margin = Inches(0.5)
            sectionMain.bottom_margin = Inches(0.5)
            sectionMain.left_margin = Inches(0.5)
            sectionMain.right_margin = Inches(0.5)
            sectionMain.header_distance = Inches(0.0)
            sectionMain.footer_distance = Inches(0.0)

            return 'First Section of A4 pages size is created.'
        else:
            return 'Page size not supported.'

        #CASE NUMBER TABLE
    def tableIdentifiers(self, caseNo1, caseNo2, parcels, FIRandDate, PSandDISTT):
        
        tableIdentifier = self.document.add_table(rows=5, cols=2)
        #TABLE STYLE
        #tableIdentifier.columns[0].width = Cm(1)
        tableIdentifier.style = 'Table Grid'
        tableIdentifier.allow_autofit =False
        #Length of table is 6309360
        tableIdentifier.rows[0].cells[0].width = Mm(130)
        tableIdentifier.rows[0].cells[1].width = Mm(130)
        tableIdentifier.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tableIdentifier.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
        #TABLE VALUES
        firstRowCells = tableIdentifier.rows[0].cells
        firstRowCells[0].paragraphs[0].add_run(f'{caseNo1}',style='TableHeading')
        firstRowCells[1].paragraphs[0].add_run(f'{caseNo1}',style='TableHeading')

        #TABLE VALUES
        secondRowCells = tableIdentifier.rows[1].cells
        secondRowCells[0].paragraphs[0].add_run(f'{caseNo2}',style='TableStyleBody')
        secondRowCells[1].paragraphs[0].add_run(f'{caseNo2}',style='TableStyleBody')

        #TABLE VALUES
        thirdRowCells = tableIdentifier.rows[2].cells
        thirdRowCells[0].paragraphs[0].add_run(f'{parcels}',style='TableStyleBody')
        thirdRowCells[1].paragraphs[0].add_run(f'{parcels}',style='TableStyleBody')

        #TABLE VALUES
        fourthRowCells = tableIdentifier.rows[3].cells
        fourthRowCells[0].paragraphs[0].add_run(f'{FIRandDate}',style='TableStyleBody')
        fourthRowCells[1].paragraphs[0].add_run(f'{FIRandDate}',style='TableStyleBody')

        #TABLE VALUES
        fifthRowCells = tableIdentifier.rows[4].cells
        fifthRowCells[0].paragraphs[0].add_run(f'{PSandDISTT}',style='TableStyleBody')
        fifthRowCells[1].paragraphs[0].add_run(f'{PSandDISTT}',style='TableStyleBody')

        self.document.add_paragraph(f" ", style="Bold10") 

    def save(self):
        self.document.save(os.path.join(UserPaths.userCaseWorkFolder(), 'Identifiers.docx'))

if __name__ == '__main__':
    i = Reports()
    i.PageLayout('A4')

    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")
    i.tableIdentifiers("PFSA2020-123456-FTM-123456", "PFSA2020-123456-FTM-123456", 1, "123 (XX.XX.XXXX)", "ABC&XYZ")

    i.save()
